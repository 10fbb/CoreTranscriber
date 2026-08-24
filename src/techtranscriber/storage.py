from __future__ import annotations

import json
import re
import threading
import wave
from dataclasses import asdict
from datetime import datetime
from pathlib import Path

import numpy as np

from .models import AudioPacket, TranscriptEntry


class SessionWriter:
    def __init__(
        self,
        output_root: Path,
        title: str,
        sample_rate: int = 48_000,
        session_started_at: float | None = None,
    ) -> None:
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        safe_title = _safe_name(title) or "Встреча"
        self.directory = output_root / f"{timestamp}_{safe_title}"
        self.directory.mkdir(parents=True, exist_ok=False)
        self.sample_rate = sample_rate
        self._entries: list[TranscriptEntry] = []
        self._entry_lock = threading.Lock()
        self._export_lock = threading.Lock()
        self._wave_lock = threading.Lock()
        self._wave = self._open_wave("meeting_audio.wav")
        self._session_started_at = session_started_at
        self._mix_start_frame = 0
        self._mix_buffer = np.empty(0, dtype=np.float32)
        self._source_frames = {"microphone": 0, "system": 0}
        self._closed = False

    @property
    def entries(self) -> list[TranscriptEntry]:
        with self._entry_lock:
            return list(self._entries)

    def _open_wave(self, name: str):
        handle = wave.open(str(self.directory / name), "wb")
        handle.setnchannels(1)
        handle.setsampwidth(2)
        handle.setframerate(self.sample_rate)
        return handle

    def write_audio(self, packet: AudioPacket) -> None:
        if self._closed:
            return
        values = np.asarray(packet.samples, dtype=np.float32).reshape(-1)
        if values.size == 0:
            return
        if packet.sample_rate != self.sample_rate:
            values = _resample(values, packet.sample_rate, self.sample_rate)
        duration = values.size / self.sample_rate
        with self._wave_lock:
            if self._session_started_at is None:
                self._session_started_at = packet.captured_at - duration
            end_frame = max(
                values.size,
                round((packet.captured_at - self._session_started_at) * self.sample_rate),
            )
            start_frame = max(0, end_frame - values.size)

            if start_frame < self._mix_start_frame:
                trim = self._mix_start_frame - start_frame
                if trim >= values.size:
                    return
                values = values[trim:]
                start_frame = self._mix_start_frame
            end_frame = start_frame + values.size
            required = end_frame - self._mix_start_frame
            if required > self._mix_buffer.size:
                self._mix_buffer = np.pad(
                    self._mix_buffer, (0, required - self._mix_buffer.size)
                )
            offset = start_frame - self._mix_start_frame
            self._mix_buffer[offset : offset + values.size] += values
            self._source_frames[packet.source] = max(
                self._source_frames[packet.source], end_frame
            )

            # Frames reached by both capture workers can no longer receive data.
            watermark = min(self._source_frames.values())
            self._flush_mix(watermark - self._mix_start_frame)

    def add_entry(self, entry: TranscriptEntry) -> None:
        with self._entry_lock:
            self._entries.append(entry)
        self.export_text_files()

    def replace_entries(self, entries: list[TranscriptEntry]) -> None:
        with self._entry_lock:
            previous = {entry.entry_id: entry for entry in self._entries}
            merged: list[TranscriptEntry] = []
            included_ids: set[str] = set()
            for entry in entries:
                original = previous.get(entry.entry_id)
                if original is not None:
                    if original.role_edited:
                        entry.role = original.role
                        entry.role_edited = True
                    if original.text_edited:
                        entry.text = original.text
                        entry.text_edited = True
                merged.append(entry)
                included_ids.add(entry.entry_id)
            for original in previous.values():
                if (
                    original.entry_id not in included_ids
                    and (original.role_edited or original.text_edited)
                ):
                    merged.append(original)
            merged.sort(key=lambda item: (item.start_seconds, item.entry_id))
            self._entries = merged
        self.export_text_files()
        if self._closed:
            self.export_docx()

    def rename_speaker(self, speaker_id: str, role: str) -> None:
        with self._entry_lock:
            for entry in self._entries:
                if entry.speaker_id == speaker_id:
                    entry.role = role
                    entry.role_edited = True
        self.export_text_files()

    def rename_entry(self, entry_id: str, role: str) -> None:
        with self._entry_lock:
            for entry in self._entries:
                if entry.entry_id == entry_id:
                    entry.role = role
                    entry.role_edited = True
                    break
        self.export_text_files()
        if self._closed:
            self.export_docx()

    def edit_entry_text(self, entry_id: str, text: str) -> None:
        with self._entry_lock:
            for entry in self._entries:
                if entry.entry_id == entry_id:
                    entry.text = text
                    entry.text_edited = True
                    break
        self.export_text_files()
        if self._closed:
            self.export_docx()

    def close(self) -> None:
        if self._closed:
            return
        self._closed = True
        with self._wave_lock:
            self._flush_mix(self._mix_buffer.size)
            self._wave.close()
        self.export_text_files()
        self.export_docx()

    def _flush_mix(self, frame_count: int) -> None:
        frame_count = max(0, min(int(frame_count), self._mix_buffer.size))
        if frame_count == 0:
            return
        samples = np.clip(self._mix_buffer[:frame_count], -1.0, 1.0)
        pcm = (samples * 32767).astype("<i2").tobytes()
        self._wave.writeframesraw(pcm)
        self._mix_buffer = self._mix_buffer[frame_count:].copy()
        self._mix_start_frame += frame_count

    def export_text_files(self) -> None:
        with self._export_lock:
            entries = self.entries
            txt = "\n".join(
                f"[{_real_time(item.created_at)}] {item.role}: {item.text}" for item in entries
            )
            (self.directory / "transcript.txt").write_text(txt, encoding="utf-8")

            md = "# Стенограмма\n\n" + "\n\n".join(
                f"**{_real_time(item.created_at)} · {item.role}**  \n{item.text}"
                for item in entries
            )
            (self.directory / "transcript.md").write_text(md, encoding="utf-8")

            payload = []
            for item in entries:
                row = asdict(item)
                row["created_at"] = item.created_at.isoformat()
                payload.append(row)
            (self.directory / "transcript.json").write_text(
                json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8"
            )

            srt_parts: list[str] = []
            for index, item in enumerate(entries, 1):
                end = item.start_seconds + max(item.duration_seconds, 1.0)
                srt_parts.append(
                    f"{index}\n{_srt_time(item.start_seconds)} --> {_srt_time(end)}\n"
                    f"{item.role}: {item.text}"
                )
            (self.directory / "transcript.srt").write_text(
                "\n\n".join(srt_parts), encoding="utf-8"
            )

    def export_docx(self) -> None:
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            return
        document = Document()
        document.add_heading("Стенограмма встречи", level=0)
        for item in self.entries:
            paragraph = document.add_paragraph()
            label = paragraph.add_run(f"{_real_time(item.created_at)} · {item.role}\n")
            label.bold = True
            label.font.size = Pt(10)
            paragraph.add_run(item.text)
        document.save(self.directory / "transcript.docx")


def _safe_name(value: str) -> str:
    value = re.sub(r"[<>:\"/\\|?*]+", "_", value.strip())
    return re.sub(r"\s+", " ", value)[:80]


def _resample(samples: np.ndarray, source_rate: int, target_rate: int) -> np.ndarray:
    if source_rate == target_rate or samples.size == 0:
        return samples.astype(np.float32, copy=False)
    target_size = max(1, round(samples.size * target_rate / source_rate))
    source_positions = np.arange(samples.size, dtype=np.float64)
    target_positions = np.linspace(0, samples.size - 1, target_size)
    return np.interp(target_positions, source_positions, samples).astype(np.float32)


def _clock(seconds: float) -> str:
    total = max(0, int(seconds))
    return f"{total // 60:02d}:{total % 60:02d}"


def _real_time(value: datetime) -> str:
    return value.strftime("%d.%m.%Y %H:%M:%S")


def _srt_time(seconds: float) -> str:
    millis = max(0, int(seconds * 1000))
    hours, remainder = divmod(millis, 3_600_000)
    minutes, remainder = divmod(remainder, 60_000)
    secs, ms = divmod(remainder, 1000)
    return f"{hours:02d}:{minutes:02d}:{secs:02d},{ms:03d}"
